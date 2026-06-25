import json
from typing import Dict, List, Any, Set
from pathlib import Path
import PyPDF2
import docx
import pandas as pd
from jinja2 import Template
import json_repair
import requests
import tempfile
import os
import re
import base64
import io
import time
import logging
from urllib.parse import urlparse
from chatbot.llm_models.llm_script import handle_bedrock_model
from chatbot.models import FileTypeChoices

# Set up logging
logger = logging.getLogger('django')
MAX_DEPTH = 1

# Additional imports for enhanced features
try:
    import fitz  # PyMuPDF for better PDF handling

    HAS_PYMUPDF = True
    logger.info("PyMuPDF available for enhanced PDF processing")
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF not available. Using PyPDF2 for text extraction only.")

try:
    from PIL import Image

    HAS_PIL = True
    logger.info("PIL available for image processing")
except ImportError:
    HAS_PIL = False
    logger.warning("PIL not available. Image processing will be limited.")

try:
    import pytesseract
    from PIL import Image as PILImage

    HAS_OCR = True
    logger.info("pytesseract available for OCR processing")
except ImportError:
    HAS_OCR = False
    logger.warning("pytesseract not available. Scanned document processing will be limited.")

try:
    import pdfplumber

    HAS_PDFPLUMBER = True
    logger.info("pdfplumber available for enhanced PDF text extraction")
except ImportError:
    HAS_PDFPLUMBER = False
    logger.warning("pdfplumber not available. Using PyMuPDF/PyPDF2 fallback.")

try:
    import openpyxl

    HAS_OPENPYXL = True
    logger.info("openpyxl available for enhanced Excel processing")
except ImportError:
    HAS_OPENPYXL = False
    logger.warning("openpyxl not available. Excel hyperlink extraction will be limited.")


class DocumentExtractor:
    """Extract structured content from documents using AWS Bedrock Llama model with enhanced features

    Usage examples:
        # Default configuration (with image extraction)
        extractor = DocumentExtractor()

        # Disable image extraction for faster processing
        extractor = DocumentExtractor(extract_images=False)

        # Custom configuration
        extractor = DocumentExtractor(
            extract_images=False,         # No image extraction
            main_doc_max_chars=15000,    # 15k chars for main doc
            subdoc_max_chars=3000,       # 3k chars for subdocs
            excel_max_rows=30,           # Only 30 rows from Excel
            excel_max_cols=15            # Only 15 columns from Excel
        )
    """

    def __init__(
            self, max_depth: int = MAX_DEPTH, max_subdocs: int = 10, enable_ocr: bool = True,
            compress_images: bool = True, extract_images: bool = False, main_doc_max_chars: int = 3000,
            subdoc_max_chars: int = 500, excel_max_rows: int = 50, excel_max_cols: int = 20,
            max_file_size_mb: int = 50
    ):
        """Initialize with enhanced features and configurable limits"""
        self.max_depth = max_depth
        self.max_subdocs = max_subdocs
        self.processed_urls: Set[str] = set()
        self.url_cache: Dict[str, str] = {}
        self.enable_ocr = enable_ocr and HAS_OCR
        self.compress_images = compress_images
        self.extract_images = extract_images  # Control image extraction

        # Configurable text limits
        self.main_doc_max_chars = main_doc_max_chars
        self.subdoc_max_chars = subdoc_max_chars
        self.excel_max_rows = excel_max_rows
        self.excel_max_cols = excel_max_cols

        # File validation
        self.allowed_extensions = {'.pdf', '.doc', '.docx', '.txt', '.csv', '.xls', '.xlsx'}
        self.max_file_size_mb = max_file_size_mb
        self.max_file_size_bytes = self.max_file_size_mb * 1024 * 1024

    def extract_urls_from_text(self, text: str) -> List[str]:
        """Extract all URLs from text content with improved regex"""
        try:
            # Log the full text for debugging
            logger.info("=" * 80)
            logger.info("EXTRACTING URLs FROM TEXT")
            logger.info("=" * 80)

            # First, let's look specifically for patterns like "word: URL" on separate lines
            lines = text.split('\n')
            manual_urls = []

            for i, line in enumerate(lines):
                line = line.strip()

                # Check if line contains http anywhere
                if 'http' in line:
                    # Extract all URLs from this line
                    url_pattern = r'https?://[^\s\n\r]+'
                    found_urls = re.findall(url_pattern, line, re.IGNORECASE)
                    manual_urls.extend(found_urls)

                # Also check if previous line ends with description and this line is a URL
                if i > 0 and line.startswith('http'):
                    if line not in manual_urls:
                        manual_urls.append(line)

            # Enhanced URL patterns for more thorough extraction
            url_patterns = [
                # Catch ALL URLs starting with http/https
                r'https?://[^\s\n\r]+',
                # Google specific patterns
                r'https://docs\.google\.com/[^/\s]+/d/[A-Za-z0-9_-]+[^\s\n\r]*',
                r'https://drive\.google\.com/[^/\s]+/d/[A-Za-z0-9_-]+[^\s\n\r]*',
            ]

            urls = []

            # Add manually found URLs first
            urls.extend(manual_urls)

            # Then use regex patterns on the full text
            for pattern in url_patterns:
                found_urls = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
                urls.extend(found_urls)

            # Clean and process URLs
            processed_urls = []
            for url in urls:
                url = url.strip()
                # Remove trailing punctuation and special chars
                url = re.sub(r'[.,;:!?)\]}>]+$', '', url)
                if url.startswith('www.'):
                    url = 'https://' + url
                processed_urls.append(url)

            # Remove duplicates while preserving order
            unique_urls = []
            seen = set()

            for url in processed_urls:
                # Normalize by removing trailing slashes
                normalized = url.rstrip('/')

                # For Google Docs/Sheets, normalize the gid parameter
                if 'docs.google.com/spreadsheets' in normalized and '#gid=' in normalized:
                    base_url = normalized.split('#gid=')[0]
                    gid_part = '#gid=' + normalized.split('#gid=')[1].split('&')[0].split('/')[0]
                    normalized = base_url + gid_part

                if normalized not in seen and len(normalized) > 10:
                    unique_urls.append(url)
                    seen.add(normalized)

            logger.info("=" * 80)
            logger.info(f"EXTRACTED {len(unique_urls)} UNIQUE URLs:")
            logger.info("=" * 80)
            for i, url in enumerate(unique_urls):
                logger.info(f"URL {i + 1}: {url}")
            logger.info("=" * 80)

            return unique_urls

        except Exception as e:
            logger.error(f"Error extracting URLs: {e}")
            return []

    def is_document_url(self, url: str, depth: int = 0) -> bool:
        """Check if URL points to a document - validates against supported formats"""
        try:
            # Exclude non-document domains/patterns
            excluded_domains = [
                'googleusercontent.com', 'gstatic.com', 'chrome.google.com',
                'googleapis.com', 'youtube.com', 'twitter.com', 'forms.google.com'
            ]

            # Check domain exclusions
            for domain in excluded_domains:
                if domain in url.lower():
                    return False

            # Special handling for Google Docs
            if any(pattern in url for pattern in [
                'docs.google.com/document',
                'drive.google.com/file',
                'docs.google.com/spreadsheets',
                'docs.google.com/forms',
                'docs.google.com/presentation'
            ]):
                return True

            parsed_url = urlparse(url)
            path = parsed_url.path.lower()

            if '.' in path:
                extension = path.rsplit('.', 1)[-1]
                # Use FileTypeChoices to validate
                return FileTypeChoices.is_valid_extension(extension)

            return False

        except Exception as e:
            logger.error(f"Error checking if URL is document: {e}")
            return False

    def convert_google_drive_url(self, url: str) -> str:
        """Convert Google URLs to downloadable formats - now includes spreadsheets and forms"""
        try:
            if 'docs.google.com/document' in url:
                if '/d/' in url:
                    doc_id = url.split('/d/')[1].split('/')[0]
                    return f"https://docs.google.com/document/d/{doc_id}/export?format=docx"

            elif 'drive.google.com/file' in url:
                if '/d/' in url:
                    file_id = url.split('/d/')[1].split('/')[0]
                    return f"https://drive.google.com/uc?id={file_id}&export=download"

            elif 'docs.google.com/spreadsheets' in url:
                if '/d/' in url:
                    sheet_id = url.split('/d/')[1].split('/')[0]
                    # Remove any gid parameter for export
                    return f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=xlsx"

            elif 'docs.google.com/forms' in url:
                # Google Forms can't be downloaded as documents
                # Return the URL as-is, it will be handled as non-downloadable
                logger.info(f"Google Form detected, cannot convert to downloadable format: {url}")
                return url

            return url
        except Exception as e:
            logger.error(f"Error converting Google Drive URL: {e}")
            return url

    def _extract_limited_excel_content(self, content_bytes: bytes, max_chars: int = None) -> str:
        """Extract limited content from Excel file for LLM processing"""
        if max_chars is None:
            max_chars = self.subdoc_max_chars

        try:
            excel_file = pd.ExcelFile(io.BytesIO(content_bytes))
            sheet_names = excel_file.sheet_names

            logger.info("=" * 80)
            logger.info(f"EXCEL FILE CONTAINS {len(sheet_names)} SHEETS:")
            for i, sheet_name in enumerate(sheet_names):
                logger.info(f"  Sheet {i + 1}: '{sheet_name}'")
            logger.info("=" * 80)

            if not sheet_names:
                return ""

            # Process sheets until we have enough content
            all_text_parts = []
            total_chars = 0
            sheets_processed = 0

            for sheet_idx, sheet_name in enumerate(sheet_names):
                if total_chars >= max_chars:
                    break

                logger.info(f"Processing sheet {sheet_idx + 1}: '{sheet_name}'")

                try:
                    # Read the sheet
                    df = pd.read_excel(
                        excel_file,
                        sheet_name=sheet_name
                    )

                    # Skip empty sheets
                    if df.empty or len(df) == 0:
                        logger.warning(f"Sheet '{sheet_name}' is empty, moving to next sheet...")
                        continue

                    # Limit rows and columns for processing
                    display_df = df.head(self.excel_max_rows)
                    if len(df.columns) > self.excel_max_cols:
                        display_df = display_df.iloc[:, :self.excel_max_cols]

                    # Convert to CSV-like format
                    csv_string = display_df.to_csv(index=False)

                    # Add sheet header if we're processing multiple sheets
                    if sheets_processed > 0:
                        all_text_parts.append(f"\n\n--- Sheet: '{sheet_name}' ---\n")

                    all_text_parts.append(csv_string)
                    sheets_processed += 1

                    # Update total characters
                    current_text = '\n'.join(all_text_parts)
                    total_chars = len(current_text)

                    logger.info(f"Sheet '{sheet_name}' added {len(csv_string)} chars (total: {total_chars} chars)")

                    # Add truncation note for this sheet if needed
                    if len(df) > self.excel_max_rows or len(df.columns) > self.excel_max_cols:
                        all_text_parts.append(
                            f"\n[Sheet '{sheet_name}': Showing {min(len(df), self.excel_max_rows)} of {len(df)} rows, "
                            f"{min(len(df.columns), self.excel_max_cols)} of {len(df.columns)} columns]"
                        )

                except Exception as e:
                    logger.error(f"Error processing sheet '{sheet_name}': {e}")
                    continue

            # If no sheets had data
            if sheets_processed == 0:
                logger.warning("All sheets are empty!")
                return "All Excel sheets are empty (no data found)"

            # Join all parts
            full_text = '\n'.join(all_text_parts)
            original_length = len(full_text)

            logger.info(f"Processed {sheets_processed} sheets with data, extracted {original_length} chars")

            # Log the content
            logger.info("=" * 80)
            logger.info("EXCEL CONTENT BEING SENT TO LLM:")
            logger.info("=" * 80)
            logger.info(full_text)
            logger.info("=" * 80)

            # Apply final character limit if needed
            if len(full_text) > max_chars:
                # Try to cut at a row boundary
                lines = full_text.split('\n')
                truncated_text = []
                current_length = 0

                for line in lines:
                    if current_length + len(line) + 1 > max_chars - 50:
                        break
                    truncated_text.append(line)
                    current_length += len(line) + 1

                full_text = '\n'.join(truncated_text) + "\n[Content truncated]"
                logger.info(f"Excel content truncated from {original_length} to {len(full_text)} chars")

            return full_text

        except Exception as e:
            logger.error(f"Error extracting Excel content: {e}")
            return ""

    def _extract_comprehensive_excel_content_for_urls(self, content_bytes: bytes) -> tuple[str, List[str]]:
        """Extract COMPLETE Excel content from ALL sheets and hyperlinks using openpyxl
        Returns: (comprehensive_text, extracted_urls)
        """
        try:
            logger.info("=" * 80)
            logger.info("EXTRACTING COMPREHENSIVE EXCEL CONTENT FOR URL EXTRACTION (OPENPYXL)")
            logger.info("=" * 80)

            if not HAS_OPENPYXL:
                logger.warning("openpyxl not available, falling back to pandas method")
                return self._extract_full_excel_content_for_urls(content_bytes), []

            wb = openpyxl.load_workbook(io.BytesIO(content_bytes), data_only=True)
            sheet_names = wb.sheetnames

            logger.info(f"Processing ALL {len(sheet_names)} sheets for content and URL extraction:")
            for i, sheet_name in enumerate(sheet_names):
                logger.info(f"  Sheet {i + 1}: '{sheet_name}'")

            all_text_parts = []
            extracted_urls = []
            total_urls_found = 0

            # Process ALL sheets without any limits
            for sheet_idx, sheet_name in enumerate(sheet_names):
                logger.info(
                    f"Processing sheet {sheet_idx + 1}/{len(sheet_names)}: '{sheet_name}' for content and URLs...")

                try:
                    sheet = wb[sheet_name]

                    # Check if sheet has data
                    if sheet.max_row == 1 and sheet.max_column == 1 and sheet.cell(1, 1).value is None:
                        logger.info(f"  Sheet '{sheet_name}' is empty, skipping...")
                        continue

                    logger.info(f"  Sheet '{sheet_name}': {sheet.max_row} rows x {sheet.max_column} columns")

                    # Extract content in multiple formats
                    sheet_text_parts = []
                    sheet_urls = []

                    # Add sheet header
                    sheet_text_parts.append(f"\n=== SHEET: {sheet_name} ===")

                    # Extract column headers (first row)
                    headers = []
                    for col in range(1, sheet.max_column + 1):
                        cell = sheet.cell(1, col)
                        header_value = cell.value
                        if header_value is not None:
                            headers.append(str(header_value))
                        else:
                            headers.append(f"Unnamed: {col - 1}")

                    sheet_text_parts.append("COLUMNS: " + " | ".join(headers))

                    # Process each row
                    for row_idx in range(1, sheet.max_row + 1):
                        row_content = []
                        row_has_content = False

                        for col_idx in range(1, sheet.max_column + 1):
                            cell = sheet.cell(row_idx, col_idx)

                            # Extract hyperlinks
                            if cell.hyperlink and cell.hyperlink.target:
                                url = cell.hyperlink.target
                                if url not in sheet_urls:
                                    sheet_urls.append(url)
                                    extracted_urls.append(url)

                            # Extract cell content
                            cell_value = cell.value
                            if cell_value is not None:
                                cell_str = str(cell_value).strip()
                                if cell_str:
                                    col_name = headers[col_idx - 1] if col_idx - 1 < len(headers) else f"Col{col_idx}"
                                    row_content.append(f"{col_name}: {cell_str}")
                                    row_has_content = True

                        if row_has_content:
                            sheet_text_parts.append(f"ROW {row_idx}: " + " | ".join(row_content))

                    # Also add CSV-like format for compatibility
                    sheet_text_parts.append("\n--- CSV FORMAT ---")
                    csv_rows = []
                    for row_idx in range(1, sheet.max_row + 1):
                        csv_row = []
                        for col_idx in range(1, sheet.max_column + 1):
                            cell = sheet.cell(row_idx, col_idx)
                            cell_value = cell.value
                            if cell_value is not None:
                                csv_row.append(str(cell_value))
                            else:
                                csv_row.append("")
                        csv_rows.append(",".join(f'"{item}"' for item in csv_row))

                    sheet_text_parts.extend(csv_rows)

                    # Join all parts for this sheet
                    sheet_content = '\n'.join(sheet_text_parts)

                    # Count URLs in this sheet for logging
                    sheet_url_count = len(sheet_urls)
                    total_urls_found += sheet_url_count

                    logger.info(
                        f"  Sheet '{sheet_name}' content: {len(sheet_content)} chars, {sheet_url_count} hyperlinks extracted")

                    all_text_parts.append(sheet_content)

                except Exception as e:
                    logger.error(f"Error processing sheet '{sheet_name}' with openpyxl: {e}")
                    continue

            # Join all sheet content
            complete_content = '\n\n'.join(all_text_parts)

            logger.info("=" * 80)
            logger.info(f"COMPREHENSIVE EXCEL EXTRACTION COMPLETE (OPENPYXL):")
            logger.info(f"  - Processed {len(sheet_names)} sheets")
            logger.info(f"  - Total content: {len(complete_content)} characters")
            logger.info(f"  - Hyperlinks extracted: {len(extracted_urls)}")
            logger.info(f"  - Total URLs found: {total_urls_found}")
            logger.info("=" * 80)

            # Log extracted URLs
            if extracted_urls:
                logger.info("EXTRACTED HYPERLINKS:")
                for i, url in enumerate(extracted_urls[:10]):  # Log first 10
                    logger.info(f"  URL {i + 1}: {url}")
                if len(extracted_urls) > 10:
                    logger.info(f"  ... and {len(extracted_urls) - 10} more URLs")

            # Log sample content
            sample_content = complete_content[:2000] if len(complete_content) > 2000 else complete_content
            logger.info("SAMPLE OF COMPREHENSIVE EXCEL CONTENT:")
            logger.info(sample_content)
            if len(complete_content) > 2000:
                logger.info(f"... [TRUNCATED - FULL CONTENT IS {len(complete_content)} CHARS] ...")
            logger.info("=" * 80)

            return complete_content, extracted_urls

        except Exception as e:
            logger.error(f"Error extracting comprehensive Excel content with openpyxl: {e}")
            # Fallback to pandas method
            return self._extract_full_excel_content_for_urls(content_bytes), []

    def _extract_full_excel_content_for_urls(self, content_bytes: bytes) -> str:
        """Fallback: Extract full Excel content specifically for URL extraction - no limits"""
        try:
            excel_file = pd.ExcelFile(io.BytesIO(content_bytes))
            sheet_names = excel_file.sheet_names

            all_text_parts = []

            # Process ALL sheets without limits
            for sheet_name in sheet_names:
                try:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    if not df.empty:
                        # Convert entire sheet to string
                        csv_string = df.to_csv(index=False)
                        all_text_parts.append(f"\n--- Sheet: '{sheet_name}' ---\n")
                        all_text_parts.append(csv_string)
                except Exception as e:
                    logger.error(f"Error processing sheet '{sheet_name}': {e}")
                    continue

            return '\n'.join(all_text_parts)

        except Exception as e:
            logger.error(f"Error extracting full Excel content: {e}")
            return ""

    def _extract_comprehensive_docx_content_for_urls(self, content_bytes: bytes) -> tuple[str, List[str]]:
        """Extract comprehensive DOCX content and hyperlinks
        Returns: (comprehensive_text, extracted_hyperlinks)
        """
        try:
            logger.info("=" * 80)
            logger.info("EXTRACTING COMPREHENSIVE DOCX CONTENT FOR URL EXTRACTION")
            logger.info("=" * 80)

            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(content_bytes)
                temp_file_path = temp_file.name

            try:
                doc = docx.Document(temp_file_path)

                # Extract all text content
                text_parts = []
                extracted_hyperlinks = []

                # Method 1: Extract from all relationships (most reliable)
                logger.info("Extracting hyperlinks from document relationships...")
                for rel_id, rel in doc.part.rels.items():
                    if hasattr(rel, 'target_ref') and rel.target_ref and rel.target_ref.startswith('http'):
                        if rel.target_ref not in extracted_hyperlinks:
                            extracted_hyperlinks.append(rel.target_ref)
                            logger.info(f"Found relationship hyperlink: {rel.target_ref}")

                # Method 2: Process paragraphs and extract hyperlinks from runs
                logger.info("Processing paragraphs for content and hyperlinks...")
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)

                    # Extract hyperlinks from paragraph runs
                    for run in para.runs:
                        if hasattr(run, '_element'):
                            # Look for hyperlink elements in the XML
                            try:
                                hyperlinks = run._element.xpath('.//w:hyperlink',
                                                                namespaces={
                                                                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                for hyperlink in hyperlinks:
                                    r_id = hyperlink.get(
                                        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                    if r_id and r_id in doc.part.rels:
                                        try:
                                            rel = doc.part.rels[r_id]
                                            if hasattr(rel,
                                                       'target_ref') and rel.target_ref and rel.target_ref not in extracted_hyperlinks:
                                                extracted_hyperlinks.append(rel.target_ref)
                                                logger.info(f"Found paragraph hyperlink: {rel.target_ref}")
                                        except:
                                            continue
                            except Exception as e:
                                logger.debug(f"Error extracting hyperlinks from run: {e}")
                                continue

                # Method 3: Process tables
                logger.info("Processing tables for content and hyperlinks...")
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(f"[Table Cell]: {cell.text}")

                            # Extract hyperlinks from table cells
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if hasattr(run, '_element'):
                                        try:
                                            hyperlinks = run._element.xpath('.//w:hyperlink',
                                                                            namespaces={
                                                                                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                            for hyperlink in hyperlinks:
                                                r_id = hyperlink.get(
                                                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                                if r_id and r_id in doc.part.rels:
                                                    try:
                                                        rel = doc.part.rels[r_id]
                                                        if hasattr(rel,
                                                                   'target_ref') and rel.target_ref and rel.target_ref not in extracted_hyperlinks:
                                                            extracted_hyperlinks.append(rel.target_ref)
                                                            logger.info(f"Found table hyperlink: {rel.target_ref}")
                                                    except:
                                                        continue
                                        except Exception as e:
                                            logger.debug(f"Error extracting hyperlinks from table cell: {e}")
                                            continue

                comprehensive_text = '\n'.join(text_parts)

                logger.info(f"DOCX extraction complete:")
                logger.info(f"  - Text content: {len(comprehensive_text)} characters")
                logger.info(f"  - Hyperlinks extracted: {len(extracted_hyperlinks)}")

                if extracted_hyperlinks:
                    logger.info("EXTRACTED HYPERLINKS:")
                    for i, url in enumerate(extracted_hyperlinks[:10]):
                        logger.info(f"  URL {i + 1}: {url}")
                    if len(extracted_hyperlinks) > 10:
                        logger.info(f"  ... and {len(extracted_hyperlinks) - 10} more URLs")

                return comprehensive_text, extracted_hyperlinks

            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

        except Exception as e:
            logger.error(f"Error extracting comprehensive DOCX content: {e}")
            # Fallback to basic text extraction
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    temp_file.write(content_bytes)
                    temp_file_path = temp_file.name

                try:
                    doc = docx.Document(temp_file_path)
                    text_parts = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)
                    return '\n'.join(text_parts), []
                finally:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)
            except:
                return "", []

    def _extract_comprehensive_pdf_content_for_urls(self, content_bytes: bytes) -> tuple[str, List[str]]:
        """Extract comprehensive PDF content and hyperlinks using PyMuPDF
        Returns: (comprehensive_text, extracted_hyperlinks)
        """
        try:
            if not HAS_PYMUPDF:
                # Fallback to basic text extraction
                text = self._extract_pdf_text_enhanced(content_bytes)
                return text, []

            logger.info("=" * 80)
            logger.info("EXTRACTING COMPREHENSIVE PDF CONTENT FOR URL EXTRACTION")
            logger.info("=" * 80)

            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(content_bytes)
                temp_file_path = temp_file.name

            try:
                doc = fitz.open(temp_file_path)
                text_parts = []
                extracted_hyperlinks = []

                logger.info(f"Processing {len(doc)} pages for content and hyperlinks...")

                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)

                    # Extract text
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

                    # Extract links/annotations
                    links = page.get_links()
                    page_hyperlinks = []

                    for link in links:
                        if 'uri' in link and link['uri']:
                            url = link['uri']
                            if url.startswith('http') and url not in extracted_hyperlinks:
                                extracted_hyperlinks.append(url)
                                page_hyperlinks.append(url)

                    if page_hyperlinks:
                        logger.info(f"  Page {page_num + 1}: {len(page_hyperlinks)} hyperlinks found")
                        for url in page_hyperlinks:
                            logger.info(f"    - {url}")

                doc.close()

                comprehensive_text = '\n'.join(text_parts)

                logger.info(f"PDF extraction complete:")
                logger.info(f"  - Text content: {len(comprehensive_text)} characters")
                logger.info(f"  - Hyperlinks extracted: {len(extracted_hyperlinks)}")

                if extracted_hyperlinks:
                    logger.info("EXTRACTED HYPERLINKS:")
                    for i, url in enumerate(extracted_hyperlinks[:10]):
                        logger.info(f"  URL {i + 1}: {url}")
                    if len(extracted_hyperlinks) > 10:
                        logger.info(f"  ... and {len(extracted_hyperlinks) - 10} more URLs")

                return comprehensive_text, extracted_hyperlinks

            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

        except Exception as e:
            logger.error(f"Error extracting comprehensive PDF content: {e}")
            # Fallback to basic text extraction
            text = self._extract_pdf_text_enhanced(content_bytes)
            return text, []

    def _extract_comprehensive_csv_content_for_urls(self, content_bytes: bytes) -> tuple[str, List[str]]:
        """Extract comprehensive CSV content (CSV files don't have hyperlinks, but we maintain consistency)
        Returns: (comprehensive_text, extracted_hyperlinks)
        """
        try:
            logger.info("=" * 80)
            logger.info("EXTRACTING COMPREHENSIVE CSV CONTENT FOR URL EXTRACTION")
            logger.info("=" * 80)

            # For CSV, there are no embedded hyperlinks, so just extract all text
            df_full = pd.read_csv(io.BytesIO(content_bytes))
            comprehensive_text = df_full.to_csv(index=False)

            logger.info(f"CSV extraction complete:")
            logger.info(f"  - Text content: {len(comprehensive_text)} characters")
            logger.info(f"  - Rows: {len(df_full)}, Columns: {len(df_full.columns)}")
            logger.info("  - No hyperlinks (CSV format doesn't support embedded links)")

            return comprehensive_text, []

        except Exception as e:
            logger.error(f"Error extracting comprehensive CSV content: {e}")
            return "", []

    def _extract_comprehensive_txt_content_for_urls(self, content_bytes: bytes) -> tuple[str, List[str]]:
        """Extract comprehensive TXT content (TXT files don't have hyperlinks, but we maintain consistency)
        Returns: (comprehensive_text, extracted_hyperlinks)
        """
        try:
            logger.info("=" * 80)
            logger.info("EXTRACTING COMPREHENSIVE TXT CONTENT FOR URL EXTRACTION")
            logger.info("=" * 80)

            # For TXT, there are no embedded hyperlinks, so just extract all text
            try:
                comprehensive_text = content_bytes.decode('utf-8', errors='ignore')
            except:
                comprehensive_text = str(content_bytes, errors='ignore')

            logger.info(f"TXT extraction complete:")
            logger.info(f"  - Text content: {len(comprehensive_text)} characters")
            logger.info("  - No hyperlinks (TXT format doesn't support embedded links)")

            return comprehensive_text, []

        except Exception as e:
            logger.error(f"Error extracting comprehensive TXT content: {e}")
            return "", []

    def _image_to_base64(self, image_bytes: bytes, image_format: str = "PNG") -> str:
        """Convert image bytes to base64 string"""
        try:
            if HAS_PIL and self.compress_images:
                # Use PIL to potentially optimize/convert image
                image = Image.open(io.BytesIO(image_bytes))
                buffer = io.BytesIO()

                # Convert to RGB if necessary
                if image.mode in ('RGBA', 'LA'):
                    background = Image.new('RGB', image.size, (255, 255, 255))
                    background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                    image = background

                # Resize if too large
                max_dimension = 1024
                if max(image.size) > max_dimension:
                    image.thumbnail((max_dimension, max_dimension), Image.Resampling.LANCZOS)

                image.save(buffer, format="JPEG", quality=85, optimize=True)
                image_bytes = buffer.getvalue()

            # Encode to base64
            base64_string = base64.b64encode(image_bytes).decode('utf-8')
            mime_type = f"image/{image_format.lower()}"
            return f"data:{mime_type};base64,{base64_string}"

        except Exception as e:
            logger.error(f"Error converting image to base64: {e}")
            return ""

    def _extract_images_from_pdf_pymupdf(self, content_bytes: bytes) -> List[Dict[str, Any]]:
        """Extract images from PDF using PyMuPDF"""
        images = []

        # Check if image extraction is enabled
        if not self.extract_images:
            return images

        if not HAS_PYMUPDF:
            return images

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(content_bytes)
                temp_file_path = temp_file.name

            try:
                doc = fitz.open(temp_file_path)

                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    image_list = page.get_images()

                    max_images_per_page = 10

                    for img_index, img in enumerate(image_list[:max_images_per_page]):
                        try:
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)

                            # Skip very small images
                            if pix.width < 100 or pix.height < 100:
                                pix = None
                                continue

                            # Skip very large images
                            if pix.width * pix.height > 2048 * 2048:
                                pix = None
                                continue

                            # Convert to PNG bytes
                            if pix.n - pix.alpha < 4:  # GRAY or RGB
                                img_bytes = pix.tobytes("png")
                                base64_image = self._image_to_base64(img_bytes, "PNG")

                                if base64_image:
                                    images.append({
                                        "page": page_num + 1,
                                        "index": img_index,
                                        "width": pix.width,
                                        "height": pix.height,
                                        "base64": base64_image,
                                        "format": "png"
                                    })

                            pix = None

                        except Exception as e:
                            logger.error(f"Error extracting image {img_index} from page {page_num + 1}: {e}")

                    if len(images) > 50:
                        logger.warning("Reached maximum image limit (50)")
                        break

                doc.close()

            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

            logger.info(f"Extracted {len(images)} images from PDF")
            return images

        except Exception as e:
            logger.error(f"Error extracting images from PDF: {e}")
            return []

    def _extract_images_from_docx(self, content_bytes: bytes) -> List[Dict[str, Any]]:
        """Extract images from DOCX file"""
        images = []

        # Check if image extraction is enabled
        if not self.extract_images:
            return images

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(content_bytes)
                temp_file_path = temp_file.name

            try:
                doc = docx.Document(temp_file_path)

                image_count = 0
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_part = rel.target_part
                            image_bytes = image_part.blob

                            if len(image_bytes) < 1000:
                                continue

                            content_type = image_part.content_type
                            image_format = "PNG"
                            if "jpeg" in content_type or "jpg" in content_type:
                                image_format = "JPEG"

                            base64_image = self._image_to_base64(image_bytes, image_format)

                            if base64_image:
                                images.append({
                                    "index": image_count,
                                    "base64": base64_image,
                                    "format": image_format.lower()
                                })
                                image_count += 1

                        except Exception as e:
                            logger.error(f"Error extracting image from DOCX: {e}")

            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

            logger.info(f"Extracted {len(images)} images from DOCX")
            return images

        except Exception as e:
            logger.error(f"Error extracting images from DOCX: {e}")
            return []

    def _perform_ocr_on_pdf(self, content_bytes: bytes) -> str:
        """Perform OCR on scanned PDF pages"""
        if not self.enable_ocr or not HAS_PYMUPDF:
            return ""

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(content_bytes)
                temp_file_path = temp_file.name

            try:
                doc = fitz.open(temp_file_path)
                ocr_text_parts = []

                for page_num in range(min(10, len(doc))):
                    page = doc.load_page(page_num)

                    # Check if page has extractable text
                    page_text = page.get_text()
                    if len(page_text.strip()) > 50:
                        continue

                    # Convert page to image for OCR
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_data = pix.tobytes("png")

                    # Perform OCR
                    image = PILImage.open(io.BytesIO(img_data))
                    ocr_text = pytesseract.image_to_string(image, lang='eng')

                    if ocr_text.strip():
                        ocr_text_parts.append(f"[Page {page_num + 1} - OCR]\n{ocr_text}")

                    pix = None

                doc.close()
                return '\n\n'.join(ocr_text_parts)

            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

        except Exception as e:
            logger.error(f"Error performing OCR: {e}")
            return ""

    def _extract_pdf_text_enhanced(self, content_bytes: bytes) -> str:
        """Enhanced PDF text extraction with multiple methods"""
        text = ""

        # Try pdfplumber first if available
        if HAS_PDFPLUMBER:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                    temp_file.write(content_bytes)
                    temp_file_path = temp_file.name

                try:
                    text_parts = []
                    import pdfplumber
                    with pdfplumber.open(temp_file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_parts.append(page_text)
                    text = '\n'.join(text_parts)
                finally:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)

                if text and len(text.strip()) > 50:
                    logger.info(f"pdfplumber extracted {len(text)} characters")
                    return text
            except Exception as e:
                logger.error(f"pdfplumber failed: {e}")

        # Try PyMuPDF next
        if HAS_PYMUPDF and not text:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                    temp_file.write(content_bytes)
                    temp_file_path = temp_file.name

                try:
                    doc = fitz.open(temp_file_path)
                    text_parts = []
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        text_parts.append(page.get_text())
                    text = '\n'.join(text_parts)
                    doc.close()
                finally:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)

                if text and len(text.strip()) > 50:
                    logger.info(f"PyMuPDF extracted {len(text)} characters")
                    return text
            except Exception as e:
                logger.error(f"PyMuPDF failed: {e}")

        # Fallback to PyPDF2
        if not text:
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
                text_parts = []
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text())
                text = '\n'.join(text_parts)
                logger.info(f"PyPDF2 extracted {len(text)} characters")
            except Exception as e:
                logger.error(f"PyPDF2 failed: {e}")

        # Try OCR if text extraction failed
        if (not text or len(text.strip()) < 50) and self.enable_ocr:
            logger.info("Attempting OCR on potentially scanned document")
            ocr_text = self._perform_ocr_on_pdf(content_bytes)
            if ocr_text:
                text = text + "\n\n[OCR Content]\n" + ocr_text if text else ocr_text

        return text

    def extract_text_from_url(self, url: str, is_subdoc: bool = False) -> tuple[
        str, List[Dict[str, Any]], Any, Dict[str, Any], str]:
        """Extract text content and images from document URL, with error handling
        Returns: (text, images, media_type, error_info, full_text_for_url_extraction)
        """
        error_info = None
        max_chars = self.subdoc_max_chars if is_subdoc else self.main_doc_max_chars

        try:
            logger.info(f"Extracting from: {url} (subdoc: {is_subdoc}, max_chars: {max_chars})")

            # Check cache
            if url in self.url_cache:
                cached_result = self.url_cache[url]
                if isinstance(cached_result, dict) and 'error' in cached_result:
                    return "", [], None, cached_result, ""
                return cached_result, [], None, None, cached_result

            # Convert Google Drive URLs to downloadable format
            download_url = self.convert_google_drive_url(url)
            if download_url is None:
                logger.info(f"Skipped non-document URL: {url}")
                return "", [], None, None, ""
            if download_url != url:
                logger.info(f"Converted to: {download_url}")

            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.5',
                'Accept-Encoding': 'gzip, deflate',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1'
            }

            response = None
            max_retries = 2
            retry_count = 0

            while retry_count < max_retries:
                try:
                    response = requests.get(download_url, headers=headers, timeout=60, allow_redirects=True)
                    response.raise_for_status()
                    if response and response.content:
                        content_size = len(response.content)
                        if content_size > self.max_file_size_bytes:
                            content_size_mb = content_size / (1024 * 1024)
                            error_info = {
                                'error': f'File size ({content_size_mb:.2f} MB) exceeds the maximum allowed '
                                         f'size of {self.max_file_size_mb} MB. Please reduce the file size.',
                                'error_type': 'file_size_exceeded',
                                'url': url
                            }
                            logger.error(f"File too large from URL {url}: {content_size_mb:.2f} MB")
                            self.url_cache[url] = error_info
                            return "", [], None, error_info, ""
                    break
                except requests.exceptions.HTTPError as e:
                    if e.response.status_code == 500 and retry_count < max_retries - 1:
                        logger.warning(f"500 error, retrying... (attempt {retry_count + 1})")
                        retry_count += 1
                        time.sleep(2)  # Wait before retry

                        # Try alternative URL format for Google Drive
                        if 'drive.google.com' in download_url and '/d/' in url:
                            file_id = url.split('/d/')[1].split('/')[0]
                            # Try alternative format
                            download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
                            logger.info(f"Trying alternative URL format: {download_url}")
                    else:
                        raise
                except requests.exceptions.Timeout:
                    if retry_count < max_retries - 1:
                        logger.warning(f"Timeout, retrying... (attempt {retry_count + 1})")
                        retry_count += 1
                        time.sleep(2)
                    else:
                        raise

            if not response:
                raise Exception("Failed to get response after retries")

            # Get content type
            content_type = response.headers.get('content-type', '').lower()
            logger.info(f"Response content_type: {content_type}")

            # Enhanced Google Drive permission detection
            if 'drive.google.com' in download_url or 'docs.google.com' in download_url:
                # Check if response is HTML-like
                if 'html' in content_type or response.text.strip().startswith(
                        '<!DOCTYPE') or response.text.strip().startswith('<html'):
                    logger.info(f"Checking for Google access denial in HTML response")
                    response_text = response.text[:1000]  # Check first 1000 chars
                    logger.info(f"Response preview: {response_text}")

                    # Check for access denial patterns
                    access_denied_patterns = [
                        'You need access',
                        'Request access',
                        'Access denied',
                        'Permission denied',
                        'Sign in to continue',
                        'This file is private',
                        'Sorry, unable to open',
                        'Google Accounts',
                        'docs.google.com/forms'
                    ]

                    for pattern in access_denied_patterns:
                        if pattern in response.text:
                            error_info = {
                                'error': f'Permission denied: Cannot access {url}',
                                'error_type': 'permission_denied',
                                'status_code': 403,
                                'url': url
                            }
                            logger.error(f"Permission denied for URL {url} - found pattern: {pattern}")
                            self.url_cache[url] = error_info
                            return "", [], None, error_info, ""

            # Validate file format based on URL extension and content
            from chatbot.models import FileTypeChoices
            from urllib.parse import urlparse
            parsed_url = urlparse(url)
            path = parsed_url.path.lower()
            url_extension = None

            # Extract extension from URL path
            if '.' in path:
                url_extension = path.rsplit('.', 1)[-1]

            # Validate extension if found in URL
            if url_extension and not FileTypeChoices.is_valid_extension(url_extension):
                error_info = {
                    'error': f'Unsupported file format: .{url_extension}',
                    'error_type': 'unsupported_format',
                    'url': url
                }
                logger.error(f"Unsupported file format .{url_extension} for URL {url}")
                self.url_cache[url] = error_info
                return "", [], None, error_info, ""

            text = ""
            images = []
            full_text_for_url_extraction = ""
            extracted_hyperlinks = []

            # Determine file type
            content_preview = response.content[:10] if response.content else b''
            is_pdf = content_preview.startswith(b'%PDF') or 'pdf' in content_type
            is_excel = any(indicator in content_type for indicator in ['spreadsheet', 'excel', 'xlsx', 'xls'])
            is_csv = 'csv' in content_type or url.lower().endswith('.csv')
            is_docx = 'word' in content_type or 'document' in content_type or 'officedocument.wordprocessing' in content_type
            is_txt = not any([is_pdf, is_excel, is_csv, is_docx])

            logger.info(f"File type - PDF: {is_pdf}, Excel: {is_excel}, CSV: {is_csv}, DOCX: {is_docx}, TXT: {is_txt}")

            # Check if it's HTML content that shouldn't be processed
            if 'html' in content_type and not any([is_pdf, is_excel, is_csv, is_docx]):
                # This is HTML content, likely an error or sign-in page
                logger.warning(f"Received HTML response for {url}, not processing as document")
                error_info = {
                    'error': 'This link returned a web page instead of a document file. This can happen with '
                             'restricted access or unsupported formats',
                    'error_type': 'invalid_content_type',
                    'url': url
                }
                self.url_cache[url] = error_info
                return "", [], None, error_info, ""

            # *** ENHANCED: Extract comprehensive content and hyperlinks for ALL formats ***
            if is_pdf:
                full_text_for_url_extraction, extracted_hyperlinks = self._extract_comprehensive_pdf_content_for_urls(
                    response.content)
                text = self._extract_pdf_text_enhanced(response.content)
                images = self._extract_images_from_pdf_pymupdf(response.content)
                media_type = FileTypeChoices.PDF
                logger.info(f"Assigned media type as: {media_type}")

            elif is_excel or 'officedocument.spreadsheet' in content_type:
                # For Excel, extract COMPREHENSIVE content for URL extraction
                logger.info("Excel file detected - extracting comprehensive content for URL detection...")
                full_text_for_url_extraction, extracted_hyperlinks = self._extract_comprehensive_excel_content_for_urls(
                    response.content)
                # Extract limited content for LLM processing
                text = self._extract_limited_excel_content(response.content, max_chars)
                media_type = FileTypeChoices.XLSX
                logger.info(f"Excel processing complete:")
                logger.info(f"  - Limited content for LLM: {len(text)} chars")
                logger.info(f"  - Comprehensive content for URLs: {len(full_text_for_url_extraction)} chars")
                logger.info(f"  - Hyperlinks extracted: {len(extracted_hyperlinks)}")

            elif is_csv:
                # For CSV, extract comprehensive content
                full_text_for_url_extraction, extracted_hyperlinks = self._extract_comprehensive_csv_content_for_urls(
                    response.content)
                # Process CSV with limited extraction for LLM
                try:
                    df = pd.read_csv(io.BytesIO(response.content), nrows=self.excel_max_rows)
                    if len(df.columns) > self.excel_max_cols:
                        df = df.iloc[:, :self.excel_max_cols]
                    text = df.to_string(max_rows=self.excel_max_rows, max_cols=self.excel_max_cols)
                    if len(text) > max_chars:
                        text = text[:max_chars] + "\n...[Content truncated]"
                    media_type = FileTypeChoices.CSV
                except Exception as e:
                    logger.error(f"Error processing CSV: {e}")
                    text = response.text[:max_chars]
                    media_type = FileTypeChoices.CSV

            elif is_docx:
                # For DOCX, extract comprehensive content and hyperlinks
                logger.info("DOCX file detected - extracting comprehensive content for URL detection...")
                full_text_for_url_extraction, extracted_hyperlinks = self._extract_comprehensive_docx_content_for_urls(
                    response.content)

                # Extract limited content for LLM processing
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    temp_file.write(response.content)
                    temp_file_path = temp_file.name

                try:
                    doc = docx.Document(temp_file_path)
                    text_parts = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)
                    text = '\n'.join(text_parts)
                    media_type = FileTypeChoices.DOCX
                    logger.info(f"Assigned media type as: {media_type}")
                finally:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)

                images = self._extract_images_from_docx(response.content)
                logger.info(f"DOCX processing complete:")
                logger.info(f"  - Limited content for LLM: {len(text)} chars")
                logger.info(f"  - Comprehensive content for URLs: {len(full_text_for_url_extraction)} chars")
                logger.info(f"  - Hyperlinks extracted: {len(extracted_hyperlinks)}")

            else:
                # For TXT and other formats
                logger.info("TXT/Other file detected - extracting content...")
                full_text_for_url_extraction, extracted_hyperlinks = self._extract_comprehensive_txt_content_for_urls(
                    response.content)
                text = response.text
                media_type = FileTypeChoices.TXT
                logger.info(f"Assigned media type as: {media_type}")

            # *** CRITICAL: Combine text-based URLs with hyperlink URLs for ALL formats ***
            combined_urls = []
            # First add hyperlink URLs (more reliable)
            combined_urls.extend(extracted_hyperlinks)
            # Then add any URLs found in text content
            text_urls = self.extract_urls_from_text(full_text_for_url_extraction)
            for url_found in text_urls:
                if url_found not in combined_urls:
                    combined_urls.append(url_found)

            # Store combined URLs for later use by appending hyperlinks to full text
            if extracted_hyperlinks:
                full_text_for_url_extraction = full_text_for_url_extraction + "\n\n=== EXTRACTED HYPERLINKS ===\n" + "\n".join(
                    extracted_hyperlinks)

            logger.info(f"URL extraction summary for {url}:")
            logger.info(f"  - Hyperlinks extracted: {len(extracted_hyperlinks)}")
            logger.info(f"  - Text-based URLs found: {len(text_urls)}")
            logger.info(f"  - Total unique URLs for processing: {len(combined_urls)}")

            # Apply character limit for subdocuments AFTER storing full text
            if is_subdoc and len(text) > max_chars:
                text = text[:max_chars] + "\n...[Content truncated]"

            # Final validation - check if we got meaningful content
            if not text or len(text.strip()) < 10:
                logger.warning(f"No meaningful content extracted from {url}")
                error_info = {
                    'error': f'No content could be extracted from {url}',
                    'error_type': 'no_content',
                    'url': url
                }
                self.url_cache[url] = error_info
                return "", [], None, error_info, ""

            if text:
                self.url_cache[url] = text

            logger.info(
                f"For url: {url}, Extracted {len(text)} characters and {len(images)} "
                f"images and file type as {media_type}"
            )

            return text, images, media_type, None, full_text_for_url_extraction

        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 403:
                error_info = {
                    'error': f'Permission denied accessing {url}',
                    'error_type': 'permission_denied',
                    'status_code': 403,
                    'url': url
                }
            elif e.response.status_code == 404:
                error_info = {
                    'error': f'Document not found at {url}',
                    'error_type': 'not_found',
                    'status_code': 404,
                    'url': url
                }
            else:
                error_info = {
                    'error': f'HTTP error {e.response.status_code} accessing {url}',
                    'error_type': 'http_error',
                    'status_code': e.response.status_code,
                    'url': url
                }
            logger.error(f"HTTP error extracting from URL {url}: {e}")
            self.url_cache[url] = error_info
            return "", [], None, error_info, ""

        except requests.exceptions.Timeout:
            error_info = {
                'error': f'Timeout accessing {url}',
                'error_type': 'timeout',
                'url': url
            }
            logger.error(f"Timeout extracting from URL {url}")
            self.url_cache[url] = error_info
            return "", [], None, error_info, ""

        except Exception as e:
            error_info = {
                'error': f'Failed to extract from {url}: {str(e)}',
                'error_type': 'extraction_error',
                'url': url
            }
            logger.error(f"Failed to extract from URL {url}: {e}")
            self.url_cache[url] = error_info
            return "", [], None, error_info, ""

    def extract_text_from_file(self, file, file_extension: str) -> tuple[str, List[Dict[str, Any]], str]:
        """Extract text content and images from various file types
        Returns: (limited_text_for_llm, images, comprehensive_text_for_urls)
        """
        try:
            file_extension = file_extension.lower().strip('.')
            text = ""
            images = []
            comprehensive_text_for_urls = ""

            # Handle file path vs file object
            if isinstance(file, (str, Path)):
                file_path = Path(file)
                if not file_path.exists():
                    raise FileNotFoundError(f"File not found: {file_path}")

                if file_extension == 'pdf':
                    with open(file_path, 'rb') as f:
                        content_bytes = f.read()
                    text = self._extract_pdf_text_enhanced(content_bytes)
                    comprehensive_text_for_urls, _ = self._extract_comprehensive_pdf_content_for_urls(content_bytes)
                    images = self._extract_images_from_pdf_pymupdf(content_bytes)
                elif file_extension in ['doc', 'docx']:
                    with open(file_path, 'rb') as f:
                        content_bytes = f.read()
                    text = self._extract_docx_text(file_path)
                    comprehensive_text_for_urls, _ = self._extract_comprehensive_docx_content_for_urls(content_bytes)
                    images = self._extract_images_from_docx(content_bytes)
                elif file_extension == 'txt':
                    with open(file_path, 'rb') as f:
                        content_bytes = f.read()
                    text = self._extract_txt_text(file_path)
                    comprehensive_text_for_urls, _ = self._extract_comprehensive_txt_content_for_urls(content_bytes)
                elif file_extension == 'csv':
                    with open(file_path, 'rb') as f:
                        content_bytes = f.read()
                    text = self._extract_csv_text(file_path)
                    comprehensive_text_for_urls, _ = self._extract_comprehensive_csv_content_for_urls(content_bytes)
                elif file_extension in ['xls', 'xlsx']:
                    with open(file_path, 'rb') as f:
                        content_bytes = f.read()
                    text = self._extract_excel_text(file_path)
                    comprehensive_text_for_urls, extracted_hyperlinks = self._extract_comprehensive_excel_content_for_urls(
                        content_bytes)
                    # Combine with hyperlinks
                    if extracted_hyperlinks:
                        comprehensive_text_for_urls = comprehensive_text_for_urls + "\n\n=== EXTRACTED HYPERLINKS ===\n" + "\n".join(
                            extracted_hyperlinks)
                    logger.info(
                        f"Main Excel file - Limited content: {len(text)} chars, Comprehensive: {len(comprehensive_text_for_urls)} chars, Hyperlinks: {len(extracted_hyperlinks)}")
                else:
                    # Default case
                    comprehensive_text_for_urls = text
            else:
                # Handle file object
                if file_extension == 'pdf':
                    file.seek(0)
                    content_bytes = file.read()
                    text = self._extract_pdf_text_enhanced(content_bytes)
                    comprehensive_text_for_urls, _ = self._extract_comprehensive_pdf_content_for_urls(content_bytes)
                    images = self._extract_images_from_pdf_pymupdf(content_bytes)
                elif file_extension in ['doc', 'docx']:
                    file.seek(0)
                    content_bytes = file.read()
                    text = self._extract_docx_text_from_object(io.BytesIO(content_bytes))
                    comprehensive_text_for_urls, _ = self._extract_comprehensive_docx_content_for_urls(content_bytes)
                    images = self._extract_images_from_docx(content_bytes)
                elif file_extension == 'txt':
                    file.seek(0)
                    content_bytes = file.read()
                    text = self._extract_txt_text_from_object(file)
                    comprehensive_text_for_urls, _ = self._extract_comprehensive_txt_content_for_urls(
                        content_bytes if isinstance(content_bytes, bytes) else content_bytes.encode('utf-8'))
                elif file_extension == 'csv':
                    file.seek(0)
                    content_bytes = file.read()
                    text = self._extract_csv_text_from_object(file)
                    comprehensive_text_for_urls, _ = self._extract_comprehensive_csv_content_for_urls(
                        content_bytes if isinstance(content_bytes, bytes) else content_bytes.encode('utf-8'))
                elif file_extension in ['xls', 'xlsx']:
                    file.seek(0)
                    content_bytes = file.read()
                    text = self._extract_excel_text_from_object(file)
                    comprehensive_text_for_urls, extracted_hyperlinks = self._extract_comprehensive_excel_content_for_urls(
                        content_bytes)
                    # Combine with hyperlinks
                    if extracted_hyperlinks:
                        comprehensive_text_for_urls = comprehensive_text_for_urls + "\n\n=== EXTRACTED HYPERLINKS ===\n" + "\n".join(
                            extracted_hyperlinks)
                    logger.info(
                        f"Main Excel file - Limited content: {len(text)} chars, Comprehensive: {len(comprehensive_text_for_urls)} chars, Hyperlinks: {len(extracted_hyperlinks)}")
                else:
                    # Default case
                    comprehensive_text_for_urls = text

            return text, images, comprehensive_text_for_urls

        except Exception as e:
            logger.error(f"Error extracting from file: {e}")
            return "", [], ""

    def _extract_pdf_text(self, file) -> str:
        """Extract text from PDF (fallback method)"""
        pdf_reader = PyPDF2.PdfReader(file)
        text_parts = []
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text_parts.append(page.extract_text())
        return '\n'.join(text_parts)

    def _extract_docx_text(self, file_path) -> str:
        """Extract text from Word document (file path)"""
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        return '\n'.join(text_parts)

    def _extract_docx_text_from_object(self, file) -> str:
        """Extract text from Word document (file object)"""
        doc = docx.Document(file)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        return '\n'.join(text_parts)

    def _extract_txt_text(self, file_path) -> str:
        """Extract text from plain text file (file path)"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _extract_txt_text_from_object(self, file) -> str:
        """Extract text from plain text file (file object)"""
        content = file.read()
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        return content

    def _extract_csv_text(self, file_path) -> str:
        """Extract text from CSV (file path) with limits"""
        df = pd.read_csv(file_path, nrows=self.excel_max_rows)
        if len(df.columns) > self.excel_max_cols:
            df = df.iloc[:, :self.excel_max_cols]
        return df.to_string(max_rows=self.excel_max_rows, max_cols=self.excel_max_cols)

    def _extract_csv_text_from_object(self, file) -> str:
        """Extract text from CSV (file object) with limits"""
        df = pd.read_csv(file, nrows=self.excel_max_rows)
        if len(df.columns) > self.excel_max_cols:
            df = df.iloc[:, :self.excel_max_cols]
        return df.to_string(max_rows=self.excel_max_rows, max_cols=self.excel_max_cols)

    def _extract_excel_text(self, file_path) -> str:
        """Extract text from Excel (file path) - first sheet only with limits"""
        excel_file = pd.ExcelFile(file_path)
        sheet_names = excel_file.sheet_names

        if not sheet_names:
            return ""

        # Only read first sheet
        df = pd.read_excel(excel_file, sheet_name=sheet_names[0], nrows=self.excel_max_rows)
        if len(df.columns) > self.excel_max_cols:
            df = df.iloc[:, :self.excel_max_cols]

        text = f"Excel file with {len(sheet_names)} sheets. Processing first sheet: '{sheet_names[0]}'\n"
        text += df.to_string(max_rows=self.excel_max_rows, max_cols=self.excel_max_cols)
        return text

    def _extract_excel_text_from_object(self, file) -> str:
        """Extract text from Excel (file object) - first sheet only with limits"""
        excel_file = pd.ExcelFile(file)
        sheet_names = excel_file.sheet_names

        if not sheet_names:
            return ""

        # Only read first sheet
        df = pd.read_excel(excel_file, sheet_name=sheet_names[0], nrows=self.excel_max_rows)
        if len(df.columns) > self.excel_max_cols:
            df = df.iloc[:, :self.excel_max_cols]

        text = f"Excel file with {len(sheet_names)} sheets. Processing first sheet: '{sheet_names[0]}'\n"
        text += df.to_string(max_rows=self.excel_max_rows, max_cols=self.excel_max_cols)
        return text

    def read_document(self, file_path: str) -> str:
        """Extract text content from various document formats"""
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower().strip('.')

        try:
            text, _, _ = self.extract_text_from_file(file_path, file_ext)
            return text
        except Exception as e:
            raise Exception(f"Error reading document: {str(e)}")

    def _find_explicit_tag_sections(self, document_text: str) -> List[str]:
        """Find explicit tag/classification sections in the document"""
        try:
            tag_sections = []

            # Common patterns for explicit tag/classification sections
            tag_patterns = [
                r'(?:tags?|keywords?|categories|classification|subject areas?|topics?|themes?):\s*([^\n\r]+)',
                r'(?:^|\n)(?:tags?|keywords?|categories|classification|subject areas?|topics?|themes?):?\s*\n([^\n\r]+(?:\n[^\n\r]+)*?)(?=\n\n|\n[A-Z]|\n\s*$|$)',
                r'(?:^|\n)(?:tags?|keywords?|categories|classification|subject areas?|topics?|themes?):?\s*\n((?:\s*[-•*]\s*[^\n\r]+\n?)+)',
                r'(?:^|\n)(?:tags?|keywords?|categories|classification|subject areas?|topics?|themes?):?\s*\n((?:\s*\d+\.\s*[^\n\r]+\n?)+)',
            ]

            doc_lower = document_text.lower()

            for pattern in tag_patterns:
                matches = re.finditer(pattern, doc_lower, re.MULTILINE | re.IGNORECASE)
                for match in matches:
                    section_content = match.group(1).strip()
                    if section_content and len(section_content) > 2:
                        tag_sections.append(section_content)

            # Remove duplicates while preserving order
            unique_sections = []
            seen_content = set()
            for section in tag_sections:
                section_key = section.lower().strip()
                if section_key not in seen_content:
                    unique_sections.append(section)
                    seen_content.add(section_key)

            return unique_sections

        except Exception as e:
            logger.error(f"Error finding explicit tag sections: {e}")
            return []

    def _validate_and_enhance_result(self, result: Dict[str, Any], document_text: str) -> Dict[str, Any]:
        """Validate and enhance the extracted result"""
        # Ensure all required fields exist
        default_response = {
            "title": "",
            "organization": "",
            "tags": [],
            "exact_content": "",
            "summary": "",
            "document_type": "",
            "key_entities": [],
            "url": [],
            "subdocument": [],
            "images": []
        }

        for key in default_response:
            if key not in result:
                result[key] = default_response[key]

        # Ensure exact_content is preserved
        if not result.get('exact_content'):
            result['exact_content'] = document_text

        # Find explicit tag sections
        explicit_tag_sections = self._find_explicit_tag_sections(document_text)

        # Validate tags
        if result.get('tags'):
            validated_tags = []
            for tag in result['tags']:
                if isinstance(tag, str):
                    # Try to parse as JSON first
                    try:
                        parsed_tag = json_repair.repair_json(tag, return_objects=True)
                        if isinstance(parsed_tag, dict) and 'text' in parsed_tag:
                            # Successfully parsed as tag dict
                            validated_tags.append(parsed_tag)
                        else:
                            # Not a valid tag dict, treat as plain string
                            validated_tags.append({"text": tag, "source": "generated"})
                    except:
                        # Failed to parse as JSON, treat as plain string
                        validated_tags.append({"text": tag, "source": "generated"})
                elif isinstance(tag, dict) and 'text' in tag:
                    # Already a proper dict with text field
                    validated_tags.append(tag)
            result['tags'] = validated_tags
        # Ensure minimum content quality
        if not result.get('title') or len(result['title'].strip()) < 3:
            content_words = document_text.split()[:10]
            result['title'] = ' '.join(content_words).strip() + '...' if content_words else 'Untitled Document'

        return result

    def extract_basic_content(
            self, document_text, company_bot, extracted_images: List[Dict[str, Any]] = None, other_data=None,
            is_subdoc=False
    ) -> Dict[str, Any]:
        """Extract basic content using Bedrock"""
        default_response = {
            "title": "",
            "organization": "",
            "tags": [],
            # "exact_content": document_text,  # Always preserve content
            "summary": "",
            "document_type": "",
            "key_entities": [],
            "url": [],
            "subdocument": [],
            "images": extracted_images or []
        }

        try:
            logger.info("Processing content with Bedrock...")
            logger.info(f"Passing Text to llm: {document_text}")

            # Preserve complete content
            complete_content = document_text

            system_prompt = [
                {
                    'text': company_bot.context
                },
            ]

            tool_context_data = json_repair.repair_json(company_bot.tool_context, return_objects=True) if isinstance(
                company_bot.tool_context, str) else company_bot.tool_context

            if isinstance(tool_context_data, list) and len(tool_context_data) > 0:
                tool_context_data = tool_context_data[0]

            end_context = company_bot.end_context

            if not end_context:
                print("Early return due to no data in end context value.")
                logger.error("Early return due to no data in end context value.")
                default_response['exact_content'] = complete_content
                return default_response

            master_document_types = None
            if company_bot and hasattr(company_bot, 'other_params') and company_bot.other_params:
                try:
                    other_params = json_repair.repair_json(company_bot.other_params, return_objects=True) if isinstance(
                        company_bot.other_params, str
                    ) else company_bot.other_params
                    master_document_types = other_params.get('master_document_types', [])
                except Exception as e:
                    print(f"Error parsing master_document_types: {e}")
                    logger.error(f"Error parsing master_document_types: {e}")
                    default_response['exact_content'] = complete_content
                    return default_response

            # Create analysis version if text is too long
            analysis_text = document_text
            max_analysis_chars = self.main_doc_max_chars

            if len(document_text) > max_analysis_chars:
                first_part = document_text[:max_analysis_chars // 2]
                last_part = document_text[-(max_analysis_chars // 2):]
                analysis_text = first_part + f"\n\n[SAMPLE - Full: {len(document_text)} chars]\n\n" + last_part

            # Include image information in context if available
            image_context = ""
            if extracted_images:
                image_context = f"\n\nDocument contains {len(extracted_images)} embedded images."
            context_data = {
                "document_text": analysis_text,
                "extracted_images": extracted_images,
                "master_tags": other_data.get('master_tag', None) if other_data else None,
                "master_document_types": master_document_types
            }
            template = Template(end_context)
            end_context = template.render(context_data)
            logger.info(f"Updated Tag Context: \n {end_context}")
            messages = [{
                'role': 'user',
                'content': [{'text': f"{end_context}"}]
            }]

            print("Bedrock: Extraction call started.")
            response = handle_bedrock_model(
                system_prompt=system_prompt,
                messages=messages,
                model_name=company_bot.llm_model,
                temperature=company_bot.bot_temperature,
                max_token=company_bot.max_token,
                company_bot=company_bot,
                tools=tool_context_data,
                aws_key=os.getenv('SG_REPO_AWS_ACCESS_KEY_ID'),
                aws_secret_key=os.getenv('SG_REPO_AWS_SECRET_ACCESS_KEY')
            )
            logger.info(f"Bedrock response type: {type(response)}")
            logger.info("Bedrock response:\n%s", json.dumps(response, indent=2))
            print(f"Bedrock response type: {type(response)}")
            print("--------\n\n")

            # *** SIMPLIFIED: Enhanced response type validation ***
            if not isinstance(response, dict):
                error_msg = ("AI processing failed - unable to extract structured data from document. "
                             "Please try uploading the file again.")
                logger.error(
                    f"LLM returned unexpected response type: {type(response)}. Expected dictionary. "
                    f"Response preview: {str(response)[:200] if response else 'None'}")

                if not is_subdoc:
                    # For main document, this is a critical error - stop processing
                    raise ValueError(error_msg)
                else:
                    # For subdocument, handle gracefully
                    default_response['exact_content'] = complete_content
                    default_response['extraction_error'] = error_msg
                    default_response['error'] = error_msg
                    default_response['error_type'] = 'ai_processing_failed'
                    default_response['title_extraction_failed'] = True
                    return default_response

            # Extract the actual data from response
            extracted_data = response.pop("parameters", response.pop("input", response))
            if not extracted_data or not isinstance(extracted_data, dict):
                error_msg = "AI processing failed - response structure is invalid. Please try uploading the file again."
                logger.error(
                    f"LLM response missing expected data structure. Response keys: {list(response.keys()) if response else 'None'}")

                if not is_subdoc:
                    raise ValueError(error_msg)

                default_response['exact_content'] = complete_content
                default_response['extraction_error'] = error_msg
                default_response['error'] = error_msg
                default_response['error_type'] = 'ai_processing_failed'
                default_response['title_extraction_failed'] = True
                return default_response

            # Preserve complete content
            extracted_data['exact_content'] = complete_content

            # Add images if available
            if extracted_images:
                extracted_data['images'] = extracted_images

            # Validate and enhance result
            result = self._validate_and_enhance_result(extracted_data, complete_content)

            if not result.get('title') or not result['title'].strip():
                error_msg = f"AI failed to extract title for {'subdocument' if is_subdoc else 'main document'}"
                logger.error(error_msg)
                if is_subdoc:
                    result['extraction_error'] = error_msg
                    result['error'] = error_msg
                    result['error_type'] = 'title_extraction_failed'
                    result['title_extraction_failed'] = True
                else:
                    # For main document, raise exception to stop processing
                    raise ValueError(error_msg)

            logger.info("Bedrock extraction successful")
            return result

        except ValueError as ve:
            # Re-raise ValueError for main document processing failures
            logger.error(f"LLM processing validation error: {str(ve)}")
            raise
        except Exception as e:
            error_msg = f"LLM processing failed with unexpected error: {str(e)}"
            logger.error(error_msg)
            default_response['exact_content'] = document_text
            if not is_subdoc:
                # For main document, raise the exception to stop processing
                raise ValueError(error_msg)
            else:
                # For subdocument, handle gracefully
                default_response['extraction_error'] = error_msg
                default_response['error'] = error_msg
                default_response['error_type'] = 'ai_processing_failed'
                default_response['title_extraction_failed'] = True
                return default_response

    def _normalize_url_for_tracking(self, url: str) -> str:
        """Normalize URL for deduplication tracking"""
        try:
            # Remove trailing slashes
            normalized = url.rstrip('/')

            # For Google Docs/Sheets, normalize parameters
            if 'docs.google.com' in normalized:
                # Extract the document/sheet ID
                if '/d/' in normalized:
                    doc_id = normalized.split('/d/')[1].split('/')[0]

                    if 'spreadsheets' in normalized:
                        # For spreadsheets, ignore gid parameter
                        base = f"https://docs.google.com/spreadsheets/d/{doc_id}"
                    elif 'document' in normalized:
                        base = f"https://docs.google.com/document/d/{doc_id}"
                    elif 'forms' in normalized:
                        base = f"https://docs.google.com/forms/d/{doc_id}"
                    else:
                        base = normalized.split('?')[0].split('#')[0]

                    return base

            # For other URLs, remove query parameters for normalization
            return normalized.split('?')[0].split('#')[0]

        except Exception as e:
            logger.error(f"Error normalizing URL {url}: {e}")
            return url

    def process_document_with_links(
            self, text: str, company_bot, comprehensive_text: str = None, processed_urls=None,
            depth=0, max_depth=MAX_DEPTH, extracted_images: List[Dict[str, Any]] = None, other_data=None
    ) -> Dict[str, Any]:
        """Process document and extract links from linked documents with enhanced URL extraction for ALL formats"""
        if processed_urls is None:
            processed_urls = set()

        try:
            # Step 1: Extract basic content from current document using Bedrock
            logger.info(f"{'  ' * depth}Processing main document with Bedrock...")
            main_result = self.extract_basic_content(text, company_bot, extracted_images, other_data)

            # *** CRITICAL FIX: Use comprehensive text for URL extraction ***
            url_extraction_text = comprehensive_text if comprehensive_text else text

            # Step 2: Extract URLs from comprehensive document content
            logger.info(f"{'  ' * depth}Extracting URLs from main document...")
            logger.info(f"{'  ' * depth}  - Using comprehensive content: {len(url_extraction_text)} chars")
            logger.info(f"{'  ' * depth}  - Limited text for LLM: {len(text)} chars")

            urls = self.extract_urls_from_text(url_extraction_text)  # ← NOW USING COMPREHENSIVE TEXT
            main_result["url"] = urls

            # Log all extracted URLs
            logger.info(f"{'  ' * depth}Total URLs extracted from main document: {len(urls)}")

            # Step 3: Process links
            subdocuments = []
            failed_links = []

            # Filter for document URLs
            document_urls = [url for url in urls if self.is_document_url(url, depth)]
            logger.info(f"{'  ' * depth}Found {len(document_urls)} document URLs in main document")

            # Process each document URL from the main document
            for main_doc_url in document_urls:
                # Normalize URL for deduplication
                normalized_url = self._normalize_url_for_tracking(main_doc_url)

                if normalized_url in processed_urls:
                    logger.info(f"{'  ' * depth}Skipping already processed URL: {main_doc_url}")
                    continue

                logger.info(f"{'  ' * depth}Processing linked document: {main_doc_url}")
                processed_urls.add(normalized_url)

                # Extract content from this linked document
                linked_text, linked_images, linked_media_type, error_info, full_text_for_urls = self.extract_text_from_url(
                    main_doc_url, is_subdoc=True
                )

                if error_info:
                    # Enhanced error handling for different error types
                    if error_info.get('error_type') == 'unsupported_format':
                        error_info['error'] = f"Unsupported file format: {error_info['error']}"

                    failed_links.append({
                        "file_url": main_doc_url,
                        "error": error_info,
                        "source_document": "main"
                    })
                    continue

                if linked_text and len(linked_text.strip()) > 10:
                    # Check if media type was determined
                    if linked_media_type is None:
                        logger.warning(f"Could not determine valid media type for {main_doc_url}")
                        failed_links.append({
                            "file_url": main_doc_url,
                            "error": {
                                'error': 'Could not determine valid file type',
                                'error_type': 'unknown_format',
                                'url': main_doc_url
                            },
                            "source_document": "main"
                        })
                        continue

                    # *** CRITICAL: Extract URLs from the COMPREHENSIVE text for ALL file formats ***
                    logger.info(f"{'  ' * depth}Extracting URLs from linked document: {main_doc_url}")
                    logger.info(f"{'  ' * depth}  - Media type: {linked_media_type}")
                    logger.info(f"{'  ' * depth}  - Using comprehensive content: {len(full_text_for_urls)} chars")

                    # Extract URLs from the comprehensive content (now includes hyperlinks for all formats)
                    links_in_subdoc = self.extract_urls_from_text(full_text_for_urls)
                    logger.info(f"{'  ' * depth}Found {len(links_in_subdoc)} total links inside {main_doc_url}")

                    # Filter for document URLs
                    subdoc_document_urls = [url for url in links_in_subdoc if self.is_document_url(url, depth)]
                    logger.info(f"{'  ' * depth}Found {len(subdoc_document_urls)} document URLs inside {main_doc_url}")

                    # Process each document URL found within the linked document
                    subdoc_count = 0
                    for sub_url in subdoc_document_urls:
                        if subdoc_count >= self.max_subdocs:
                            logger.info(f"{'  ' * (depth + 1)}Reached max subdocs limit ({self.max_subdocs})")
                            break

                        # Normalize URL for deduplication
                        normalized_sub_url = self._normalize_url_for_tracking(sub_url)

                        if normalized_sub_url in processed_urls:
                            logger.info(f"{'  ' * (depth + 1)}Skipping already processed subdocument URL: {sub_url}")
                            continue

                        logger.info(f"{'  ' * (depth + 1)}Processing subdocument: {sub_url}")
                        processed_urls.add(normalized_sub_url)
                        subdoc_count += 1

                        # Extract content from subdocument URL
                        sub_text, sub_images, sub_media_type, sub_error_info, _ = self.extract_text_from_url(
                            sub_url, is_subdoc=True
                        )
                        logger.info(f"for url: {sub_url}, extracted sub_text is: {sub_text}")

                        if sub_error_info:
                            logger.info(f"{'  ' * (depth + 1)}Subdocument failed: {sub_error_info}")

                            # Enhance error message for unsupported formats
                            if sub_error_info.get('error_type') == 'unsupported_format':
                                sub_error_info[
                                    'error'] = f"Unsupported file format in linked document: {sub_error_info['error']}"

                            failed_links.append({
                                "file_url": sub_url,
                                "error": sub_error_info,
                                "source_document": main_doc_url
                            })
                        else:
                            # Successfully accessed - process subdocument with LLM
                            if sub_text and len(sub_text.strip()) > 10:
                                # Get the downloadable URL
                                downloadable_url = self.convert_google_drive_url(sub_url)

                                # Check if media type was determined
                                if sub_media_type is None:
                                    logger.warning(f"Could not determine valid media type for {sub_url}")
                                    failed_links.append({
                                        "file_url": sub_url,
                                        "error": {
                                            'error': 'Could not determine valid file type',
                                            'error_type': 'unknown_format',
                                            'url': sub_url
                                        },
                                        "source_document": main_doc_url
                                    })
                                    continue

                                # Process subdocument content with Bedrock
                                subdoc_result = self.extract_basic_content(
                                    sub_text,
                                    company_bot,
                                    sub_images,
                                    other_data,
                                    is_subdoc=True
                                )

                                # Check for any extraction errors in subdocument
                                if (subdoc_result.get('title_extraction_failed') or
                                        subdoc_result.get('extraction_error') or
                                        subdoc_result.get('error') or
                                        subdoc_result.get('error_type')):
                                    error_message = (subdoc_result.get('error') or
                                                     subdoc_result.get('extraction_error') or
                                                     'LLM failed to extract title from subdocument')

                                    error_type = (subdoc_result.get('error_type') or
                                                  'title_extraction_failed')

                                    logger.error(f"Subdocument extraction failed for {sub_url}: {error_message}")
                                    failed_links.append({
                                        "file_url": sub_url,
                                        "error": {
                                            'error': error_message,
                                            'error_type': error_type,
                                            'url': sub_url
                                        },
                                        "source_document": main_doc_url
                                    })
                                    continue

                                # Create subdocument entry (without "url" field)
                                subdoc_entry = {
                                    "title": subdoc_result.get(
                                        "title",
                                        f"Document from {Path(urlparse(main_doc_url).path).name or 'linked document'}"
                                    ),
                                    "file_url": downloadable_url,
                                    "media_type": sub_media_type,
                                    "source_document": main_doc_url,
                                    "exact_content": sub_text,
                                    "summary": subdoc_result.get("summary", ""),
                                    "tags": subdoc_result.get("tags", []),
                                    "organization": subdoc_result.get("organization", ""),
                                    "document_type": subdoc_result.get("document_type", ""),
                                    "key_entities": subdoc_result.get("key_entities", []),
                                    "subdocument": [],
                                    "images": sub_images or []
                                }
                                subdocuments.append(subdoc_entry)
                            else:
                                logger.warning(f"Subdocument {sub_url} has insufficient content")
                                failed_links.append({
                                    "file_url": sub_url,
                                    "error": {
                                        'error': 'Document has insufficient content (less than 10 characters)',
                                        'error_type': 'insufficient_content',
                                        'url': sub_url
                                    },
                                    "source_document": main_doc_url
                                })
                else:
                    logger.warning(f"Linked document {main_doc_url} has insufficient content: {linked_text}")
                    failed_links.append({
                        "file_url": main_doc_url,
                        "error": {
                            'error': 'Document has insufficient content (less than 10 characters)',
                            'error_type': 'insufficient_content',
                            'url': main_doc_url
                        },
                        "source_document": "main"
                    })

            main_result["subdocument"] = subdocuments
            main_result["failed_links"] = failed_links

            # Log summary
            logger.info(f"{'  ' * depth}Processing complete:")
            logger.info(f"{'  ' * depth}  - URLs in main document: {len(urls)}")
            logger.info(f"{'  ' * depth}  - Document URLs in main: {len(document_urls)}")
            logger.info(f"{'  ' * depth}  - Successfully processed subdocuments: {len(subdocuments)}")
            logger.info(f"{'  ' * depth}  - Failed: {len(failed_links)}")
            logger.info(f"{'  ' * depth}  - Total URLs processed: {len(processed_urls)}")

            return main_result

        except ValueError as ve:
            logger.error(f"Main document processing failed: {str(ve)}")
            raise
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                "title": "",
                "organization": "",
                "tags": [],
                "exact_content": text,
                "summary": "",
                "document_type": "",
                "key_entities": [],
                "url": [],
                "subdocument": [],
                "failed_links": [],
                "images": extracted_images or []
            }

    def _determine_media_type_from_url(self, url: str) -> str:
        """Determine media type from URL"""
        try:
            parsed_url = urlparse(url)
            path = parsed_url.path.lower()

            # Extract extension if present
            if '.' in path:
                extension = path.rsplit('.', 1)[-1]

                # Check if it's a valid extension first
                if not FileTypeChoices.is_valid_extension(extension):
                    logger.warning(f"Invalid extension {extension} in URL {url}")
                    return None  # Return None for invalid extensions

                # Use the existing method instead of hardcoding
                mime_type = FileTypeChoices.get_mime_from_extension(extension)
                if mime_type:
                    return mime_type.value
                else:
                    # Extension is valid but not mapped - default to TXT
                    logger.warning(f"No MIME type mapping for valid extension {extension}")
                    return FileTypeChoices.TXT.value

            # No extension found - default to TXT
            return FileTypeChoices.TXT.value

        except Exception as e:
            logger.error(f"Error determining media type from URL {url}: {e}")
            return FileTypeChoices.TXT.value

    def _get_comprehensive_content_for_url_extraction(self, document_text: str, other_data=None) -> str:
        """Get comprehensive content for URL extraction from the original file"""
        try:
            # If we have the comprehensive content stored in other_data, use it
            if other_data and 'comprehensive_text_for_urls' in other_data:
                comprehensive_text = other_data['comprehensive_text_for_urls']
                logger.info(f"Using stored comprehensive content: {len(comprehensive_text)} chars")
                return comprehensive_text

            # Fallback to the document_text if no comprehensive content available
            logger.info(f"No comprehensive content available, using document text: {len(document_text)} chars")
            return document_text

        except Exception as e:
            logger.error(f"Error getting comprehensive content: {e}")
            return document_text

    def extract_with_bedrock(
            self, document_text, company_bot, extracted_images: List[Dict[str, Any]] = None, other_data=None
    ) -> Dict[str, Any]:
        """Main entry point - processes document with recursive link extraction"""
        try:
            logger.info("Starting document processing with recursive link extraction...")

            # *** CRITICAL FIX: Get comprehensive content for URL extraction ***
            comprehensive_text_for_urls = self._get_comprehensive_content_for_url_extraction(
                document_text, other_data
            )

            result = self.process_document_with_links(
                text=document_text,  # Limited text for LLM
                comprehensive_text=comprehensive_text_for_urls,  # Full text for URL extraction
                company_bot=company_bot,
                extracted_images=extracted_images,
                other_data=other_data
            )
            return result
        except ValueError as ve:
            # Re-raise ValueError so it can be handled by the calling function
            logger.error(f"Document processing validation failed: {str(ve)}")
            raise  # This allows the error to propagate to get_doc_tags_from_ai()
        except Exception as e:
            logger.error(f"Document processing failed with unexpected error: {str(e)}")
            return {
                "title": "",
                "organization": "",
                "tags": [],
                "summary": "",
                "document_type": "",
                "key_entities": [],
                "url": [],
                "subdocument": [],
                "images": extracted_images or []
            }

    def extract_with_llm(self, text: str, company_bot=None, max_chars: int = 6000,
                         extracted_images: List[Dict[str, Any]] = None, other_data=None) -> Dict[str, Any]:
        """Extract structured information using AWS Bedrock Llama"""
        # Don't truncate - preserve complete content
        return self.extract_with_bedrock(
            document_text=text, company_bot=company_bot, extracted_images=extracted_images, other_data=other_data
        )

    def process_document_from_url(self, url: str, company_bot=None) -> Dict[str, Any]:
        """Process document directly from URL"""
        try:
            text_content, extracted_images, extracted_media_type, error_info, _ = self.extract_text_from_url(
                url, is_subdoc=False
            )

            if error_info:
                return {
                    "error": error_info['error'],
                    "error_type": error_info.get('error_type', 'unknown'),
                    "file_path": url,
                    "file_name": Path(url).name,
                }

            if not text_content or len(text_content.strip()) < 10:
                raise ValueError("Document appears to be empty or unreadable")

            extracted_info = self.extract_with_llm(text_content, company_bot, extracted_images=extracted_images)

            result = {
                "file_path": url,
                "file_name": Path(url).name,
                "text_length": len(text_content),
                **extracted_info
            }

            return result

        except Exception as e:
            return {
                "error": str(e),
                "file_path": url,
                "file_name": "Unknown",
            }

    def process_document(self, file_path: str, company_bot=None, other_data=None) -> Dict[str, Any]:
        """Process document from file path"""
        try:
            # Read document content with enhanced extraction
            text_content, extracted_images, comprehensive_text_for_urls = self.extract_text_from_file(file_path, Path(
                file_path).suffix.strip('.'))

            if not text_content or len(text_content.strip()) < 10:
                raise ValueError("Document appears to be empty or unreadable")

            # Extract structured information using LLM
            extracted_info = self.extract_with_llm(
                text=text_content,
                company_bot=company_bot,
                extracted_images=extracted_images,
                other_data=other_data,
            )

            # Add metadata
            result = {
                "file_path": str(file_path),
                "file_name": Path(file_path).name,
                "text_length": len(text_content),
                **extracted_info
            }

            return result

        except Exception as e:
            return {
                "error": str(e),
                "file_path": str(file_path),
                "file_name": Path(file_path).name if Path(file_path).exists() else "Unknown",
            }


# Additional helper functions for file object processing
def extract_tags_from_document_url(url: str, company_bot) -> Dict[str, Any]:
    """Extract structured information from document URL"""
    default_response = {
        "title": "",
        "organization": "",
        "tags": [],
        "exact_content": "",
        "summary": "",
        "document_type": "",
        "key_entities": [],
        "url": [],
        "subdocument": [],
        "images": []
    }
    try:
        # Parse other_params from company_bot
        extractor_config = {}
        if company_bot and hasattr(company_bot, 'other_params') and company_bot.other_params:
            try:
                other_params = json.loads(company_bot.other_params) if isinstance(
                    company_bot.other_params, str) else company_bot.other_params

                # Extract DocumentExtractor configuration
                extractor_config = {
                    'max_depth': other_params.get('max_depth', MAX_DEPTH),
                    'max_subdocs': other_params.get('max_subdocs', 10),
                    'enable_ocr': other_params.get('enable_ocr', True),
                    'compress_images': other_params.get('compress_images', True),
                    'extract_images': other_params.get('extract_images', False),
                    'main_doc_max_chars': other_params.get('main_doc_max_chars', 3000),
                    'subdoc_max_chars': other_params.get('subdoc_max_chars', 500),
                    'excel_max_rows': other_params.get('excel_max_rows', 50),
                    'excel_max_cols': other_params.get('excel_max_cols', 20),
                    'max_file_size_mb': other_params.get('max_file_size_mb', 50),
                }
            except Exception as e:
                logger.error(f"Error parsing other_params: {e}")

        extractor = DocumentExtractor(**extractor_config)
        result = extractor.process_document_from_url(url, company_bot)
        return result
    except Exception as e:
        logger.error(f"Error extracting tags from URL: {e}")
        return default_response


def extract_tags_from_document_file(file, company_bot, file_extension, other_data):
    """Extract structured information from document file"""
    default_response = {
        "title": "",
        "organization": "",
        "tags": [],
        "exact_content": "",
        "summary": "",
        "document_type": "",
        "key_entities": [],
        "url": [],
        "subdocument": [],
        "images": []
    }
    try:
        extractor_config = {}
        if company_bot and hasattr(company_bot, 'other_params') and company_bot.other_params:
            try:
                other_params = json.loads(company_bot.other_params) if isinstance(
                    company_bot.other_params, str
                ) else company_bot.other_params

                # Extract DocumentExtractor configuration
                extractor_config = {
                    'max_depth': other_params.get('max_depth', MAX_DEPTH),
                    'max_subdocs': other_params.get('max_subdocs', 10),
                    'enable_ocr': other_params.get('enable_ocr', True),
                    'compress_images': other_params.get('compress_images', True),
                    'extract_images': other_params.get('extract_images', False),
                    'main_doc_max_chars': other_params.get('main_doc_max_chars', 3000),
                    'subdoc_max_chars': other_params.get('subdoc_max_chars', 500),
                    'excel_max_rows': other_params.get('excel_max_rows', 50),
                    'excel_max_cols': other_params.get('excel_max_cols', 20),
                    'max_file_size_mb': other_params.get('max_file_size_mb', 50),
                }
            except Exception as e:
                logger.error(f"Error parsing other_params: {e}")

        max_file_size_mb = extractor_config.get('max_file_size_mb', 50)
        max_file_size_bytes = max_file_size_mb * 1024 * 1024
        file_size = 0
        if hasattr(file, 'size'):
            file_size = file.size
        elif hasattr(file, 'seek') and hasattr(file, 'tell'):
            current_position = file.tell()
            file.seek(0, 2)  # Seek to end
            file_size = file.tell()
            file.seek(current_position)

        if file_size > max_file_size_bytes:
            file_size_mb = file_size / (1024 * 1024)
            error_msg = (f"File size ({file_size_mb:.2f} MB) exceeds the maximum allowed size "
                         f"of {max_file_size_mb} MB. Please reduce the file size.")
            logger.error(error_msg)
            raise ValueError(error_msg)

        extractor = DocumentExtractor(**extractor_config)

        # *** CRITICAL FIX: Extract both limited and comprehensive content ***
        document_text, extracted_images, comprehensive_text_for_urls = extractor.extract_text_from_file(file, file_extension)

        if not document_text:
            return default_response

        # *** CRITICAL FIX: Pass comprehensive content in other_data ***
        if not other_data:
            other_data = {}
        other_data['comprehensive_text_for_urls'] = comprehensive_text_for_urls

        # Extract information using Bedrock with URL processing
        result = extractor.extract_with_llm(
            document_text, company_bot, extracted_images=extracted_images, other_data=other_data
        )

        return result

    except ValueError as ve:
        error_message = str(ve)
        logger.error(f"Processing error: {error_message}")

        # *** FIX: Return error response instead of default_response ***
        error_response = {
            "error": error_message,
            "title": "",
            "organization": "",
            "tags": [],
            "exact_content": "",
            "summary": "",
            "document_type": "",
            "key_entities": [],
            "url": [],
            "subdocument": [],
            "images": []
        }

        # *** NEW: Distinguish between different types of ValueError ***
        if "file size" in error_message.lower() or "exceeds the maximum allowed size" in error_message.lower():
            # File size validation error
            error_response["error_type"] = "file_size_exceeded"
        elif "llm returned" in error_message.lower() or "unexpected list format" in error_message.lower() or "plain text instead" in error_message.lower():
            # LLM response format error
            error_response["error_type"] = "llm_response_format_error"
        elif "failed to extract title" in error_message.lower():
            # Title extraction error
            error_response["error_type"] = "title_extraction_error"
        elif "processing failed with unexpected error" in error_message.lower():
            # LLM processing error
            error_response["error_type"] = "llm_processing_error"
        else:
            # Generic validation error
            error_response["error_type"] = "validation_error"

        return error_response  # ← NOW RETURNS ERROR INSTEAD OF DEFAULT

    except Exception as e:
        # *** NEW: Handle any other unexpected errors ***
        error_message = f"Unexpected error during document processing: {str(e)}"
        logger.error(error_message)
        return {
            "error": error_message,
            "error_type": "unexpected_error",
            "title": "",
            "organization": "",
            "tags": [],
            "exact_content": "",
            "summary": "",
            "document_type": "",
            "key_entities": [],
            "url": [],
            "subdocument": [],
            "images": []
        }

def get_doc_tags_from_ai(file, company_bot, file_extension, other_data):
    """Main entry point for document processing with improved error handling"""
    try:
        result = extract_tags_from_document_file(file, company_bot, file_extension, other_data)
        print("Final result: ", result)
        logger.info("Final Extraction Result:\n%s", json.dumps(result, indent=2, ensure_ascii=False))
        return result
    except ValueError as ve:
        error_message = str(ve)
        logger.error(f"Processing error: {error_message}")

        # *** SIMPLIFIED: Common error response for all AI processing failures ***
        if any(keyword in error_message.lower() for keyword in [
            "ai processing failed", "unable to extract structured data",
            "llm returned", "unexpected", "processing failed"
        ]):
            return {
                "error": "AI processing failed - unable to extract structured data from document. Please try uploading the file again.",
                "error_type": "ai_processing_failed",
                "title": "",
                "organization": "",
                "tags": [],
                "exact_content": "",
                "summary": "",
                "document_type": "",
                "key_entities": [],
                "url": [],
                "subdocument": [],
                "images": []
            }
        elif "file size" in error_message.lower() or "exceeds the maximum allowed size" in error_message.lower():
            # File size validation error
            return {
                "error": error_message,
                "error_type": "file_size_exceeded",
                "title": "",
                "organization": "",
                "tags": [],
                "exact_content": "",
                "summary": "",
                "document_type": "",
                "key_entities": [],
                "url": [],
                "subdocument": [],
                "images": []
            }
        else:
            # Generic validation error
            return {
                "error": f"Document processing failed: {error_message}",
                "error_type": "validation_error",
                "title": "",
                "organization": "",
                "tags": [],
                "exact_content": "",
                "summary": "",
                "document_type": "",
                "key_entities": [],
                "url": [],
                "subdocument": [],
                "images": []
            }
    except Exception as e:
        # *** SIMPLIFIED: Handle any other unexpected errors ***
        error_message = "AI processing failed - unable to extract structured data from document. Please try uploading the file again."
        logger.error(f"Unexpected error during document processing: {str(e)}")
        return {
            "error": error_message,
            "error_type": "ai_processing_failed",
            "title": "",
            "organization": "",
            "tags": [],
            "exact_content": "",
            "summary": "",
            "document_type": "",
            "key_entities": [],
            "url": [],
            "subdocument": [],
            "images": []
        }
