import io
import zipfile
import requests
from django.utils.text import slugify
from django.http import HttpResponseRedirect
from django.contrib import admin
from django.utils.http import urlencode
from io import BytesIO
from django.utils.timezone import localtime
from docx import Document
from django.http import HttpResponse
from django.forms.models import model_to_dict
from docx.shared import Inches
import tempfile
import os
from chatbot.models import MediaTypeChoices
from urllib.parse import urlparse


@admin.action(description='Export selected stories')
def redirect_to_export_view(modeladmin, request, queryset):
    selected = queryset.values_list('pk', flat=True)
    query_string = urlencode({'ids': ','.join(map(str, selected))})
    return HttpResponseRedirect(f'{request.path}export_stories/?{query_string}')


def get_all_other_params_keys(stories):
    keys = set()
    for story in stories:
        if isinstance(story.other_params, dict):
            keys.update(story.other_params.keys())
    return sorted(keys)


def generate_csv_response(dataset):
    response = HttpResponse(dataset.export('csv'), content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename=stories.csv'
    return response


def generate_xls_response(dataset):
    response = HttpResponse(dataset.export('xls'), content_type='application/vnd.ms-excel')
    response['Content-Disposition'] = 'attachment; filename=stories.xls'
    return response


def generate_docx_response(stories, fields_to_export):
    document = Document()
    headers = get_story_fields(stories, fields_to_export)

    for i, story in enumerate(stories, start=1):
        document.add_heading(f'Story {i}: {story.title}', level=1)

        row_data = get_story_data(story, headers)
        for field, value in zip(headers, row_data):
            if field == "story_media_urls":
                if not value:
                    continue
                document.add_paragraph(f"{field.replace('_', ' ').title()}:")
                urls = value.split(', ')
                for url in urls:
                    if not url.lower().endswith('.pdf'):
                        try:
                            img_response = requests.get(url)
                            if img_response.status_code == 200:
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp_img:
                                    tmp_img.write(img_response.content)
                                    tmp_img.flush()
                                    document.add_paragraph(url)  # Show URL
                                    try:
                                        # Try showing if it's image
                                        document.add_picture(
                                            tmp_img.name,width=Inches(2.5)
                                        )
                                    except Exception as e:
                                        document.add_paragraph(f"(Preview not available: {e})")
                        except Exception as e:
                            document.add_paragraph(f"Failed to load image: {url} ({e})")
                    else:
                        document.add_paragraph(url)  # Non-image media, just show link
            else:
                document.add_paragraph(f"{field.replace('_', ' ').title()}: {value}")

        if i != len(stories):
            document.add_page_break()

    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)

    response = HttpResponse(
        doc_io.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=stories.docx'
    return response


def get_story_fields(stories, fields_to_export):
    # base_fields = [field.name for field in Story._meta.fields if field.name != 'other_params']
    headers = fields_to_export.copy()

    extra_fields = set()

    for story in stories:
        if isinstance(story.other_params, dict):
            extra_fields.update(story.other_params.keys())

    # headers = base_fields + sorted(extra_fields)
    headers += sorted(extra_fields)
    headers.append("story_pdfs")
    headers.append("story_media_urls")

    return headers


def get_story_data(story, headers):
    data = []
    story_dict = model_to_dict(story)

    for field in headers:
        if field == 'story_pdfs':
            pdf = story.story_media.filter(media_type=MediaTypeChoices.PDF).first()
            value = pdf.get_public_url() if pdf else ''
        elif field == 'story_media_urls':
            media_urls = [
                media.get_public_url()
                for media in story.story_media.all()
                if media.media_type != MediaTypeChoices.PDF and media.get_public_url()
            ]
            if media_urls:
                value = ', '.join(media_urls)
            else:
                value = None

        elif field in story_dict:
            value = story_dict[field]
            if hasattr(value, '__str__'):
                value = str(value)
        elif field == 'created_at' and story.created_at:
            value = localtime(story.created_at).replace(tzinfo=None)
        else:
            value = story.other_params.get(field, '') if story.other_params else ''
        data.append(value)
    return data


def generate_zip_response(stories):
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        for story in stories:
            pdfs = story.story_media.filter(media_type=MediaTypeChoices.PDF)
            for i, pdf in enumerate(pdfs, start=1):
                print(f"Story {story.id} has {pdfs.count()} PDFs")
                url = pdf.get_public_url()
                if not url:
                    continue
                try:
                    response = requests.get(url)
                    if response.status_code == 200:
                        # Use pdf.name, fallback to something if it's missing
                        base_name = get_filename_from_url(url) or pdf.name or f"story_{story.id}_media_{i}"
                        print("Pdf name: ", base_name)
                        # safe_name = slugify(base_name)
                        filename = f"{base_name}_{story.id}.pdf"
                        zip_file.writestr(filename, response.content)
                    else:
                        print(f"Failed to download from {url}, status code {response.status_code}")
                except Exception as e:
                    print(f"Error downloading {url}: {e}")

    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename=stories.zip'
    return response

def get_filename_from_url(url):
    path = urlparse(url).path
    filename = os.path.basename(path)
    print("Filename from url is: ", filename)
    return filename
