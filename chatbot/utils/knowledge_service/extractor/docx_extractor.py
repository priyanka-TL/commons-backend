"""DOCX extraction functionality"""

import io
import os
import logging
import tempfile
from typing import List, Dict, Any, Tuple
import docx

logger = logging.getLogger('django')


class DOCXExtractor:
    """Handles DOCX content extraction"""

    def __init__(self, image_processor):
        self.image_processor = image_processor

    def extract_comprehensive_content_for_urls(self, content_bytes: bytes) -> Tuple[str, List[str]]:
        """Extract comprehensive DOCX content and hyperlinks

        Args:
            content_bytes: DOCX file content as bytes

        Returns:
            Tuple of (comprehensive_text, extracted_hyperlinks)
        """
        try:
            logger.info("=" * 80)
            logger.info("EXTRACTING COMPREHENSIVE DOCX CONTENT FOR URL EXTRACTION")
            logger.info("=" * 80)

            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(content_bytes)
                temp_file_path = temp_file.name

            try:
                doc = docx.Document(temp_file_path)

                # Extract all text content
                text_parts = []
                extracted_hyperlinks = []

                # Method 1: Extract from all relationships (most reliable)
                logger.info("Extracting hyperlinks from document relationships...")
                for rel_id, rel in doc.part.rels.items():
                    if hasattr(rel, 'target_ref') and rel.target_ref and rel.target_ref.startswith('http'):
                        if rel.target_ref not in extracted_hyperlinks:
                            extracted_hyperlinks.append(rel.target_ref)
                            logger.info(f"Found relationship hyperlink: {rel.target_ref}")

                # Method 2: Process paragraphs and extract hyperlinks from runs
                logger.info("Processing paragraphs for content and hyperlinks...")
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)

                    # Extract hyperlinks from paragraph runs
                    for run in para.runs:
                        if hasattr(run, '_element'):
                            # Look for hyperlink elements in the XML
                            try:
                                hyperlinks = run._element.xpath('.//w:hyperlink',
                                                                namespaces={
                                                                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                for hyperlink in hyperlinks:
                                    r_id = hyperlink.get(
                                        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                    if r_id and r_id in doc.part.rels:
                                        try:
                                            rel = doc.part.rels[r_id]
                                            if hasattr(rel,
                                                       'target_ref') and rel.target_ref and rel.target_ref not in extracted_hyperlinks:
                                                extracted_hyperlinks.append(rel.target_ref)
                                                logger.info(f"Found paragraph hyperlink: {rel.target_ref}")
                                        except:
                                            continue
                            except Exception as e:
                                logger.debug(f"Error extracting hyperlinks from run: {e}")
                                continue

                # Method 3: Process tables
                logger.info("Processing tables for content and hyperlinks...")
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(f"[Table Cell]: {cell.text}")

                            # Extract hyperlinks from table cells
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if hasattr(run, '_element'):
                                        try:
                                            hyperlinks = run._element.xpath('.//w:hyperlink',
                                                                            namespaces={
                                                                                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                            for hyperlink in hyperlinks:
                                                r_id = hyperlink.get(
                                                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                                if r_id and r_id in doc.part.rels:
                                                    try:
                                                        rel = doc.part.rels[r_id]
                                                        if hasattr(rel,
                                                                   'target_ref') and rel.target_ref and rel.target_ref not in extracted_hyperlinks:
                                                            extracted_hyperlinks.append(rel.target_ref)
                                                            logger.info(f"Found table hyperlink: {rel.target_ref}")
                                                    except:
                                                        continue
                                        except Exception as e:
                                            logger.debug(f"Error extracting hyperlinks from table cell: {e}")
                                            continue

                comprehensive_text = '\n'.join(text_parts)

                logger.info(f"DOCX extraction complete:")
                logger.info(f"  - Text content: {len(comprehensive_text)} characters")
                logger.info(f"  - Hyperlinks extracted: {len(extracted_hyperlinks)}")

                if extracted_hyperlinks:
                    logger.info("EXTRACTED HYPERLINKS:")
                    for i, url in enumerate(extracted_hyperlinks[:10]):
                        logger.info(f"  URL {i + 1}: {url}")
                    if len(extracted_hyperlinks) > 10:
                        logger.info(f"  ... and {len(extracted_hyperlinks) - 10} more URLs")

                return comprehensive_text, extracted_hyperlinks

            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

        except Exception as e:
            logger.error(f"Error extracting comprehensive DOCX content: {e}")
            # Fallback to basic text extraction
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    temp_file.write(content_bytes)
                    temp_file_path = temp_file.name

                try:
                    doc = docx.Document(temp_file_path)
                    text_parts = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)
                    return '\n'.join(text_parts), []
                finally:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)
            except:
                return "", []

    def extract_text(self, file_path) -> str:
        """Extract text from Word document (file path)

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        return '\n'.join(text_parts)

    def extract_text_from_object(self, file) -> str:
        """Extract text from Word document (file object)

        Args:
            file: DOCX file object

        Returns:
            Extracted text content
        """
        doc = docx.Document(file)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        return '\n'.join(text_parts)